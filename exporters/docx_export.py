import os
import logging
import re
from typing import Dict, Any, List, Optional
from pathlib import Path

import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

from services.style_extract import StyleProfile

logger = logging.getLogger(__name__)

class DocxExporter:
    def __init__(self):
        self.document = None
        
    def _setup_document_styles(self, style_profile: StyleProfile) -> None:
        if not self.document:
            return
            
        styles = self.document.styles
        
        try:
            heading_style = styles.add_style('CV Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_style.paragraph_format.space_after = Pt(6)
        except:
            heading_style = styles['CV Heading'] if 'CV Heading' in styles else styles['Heading 2']
        
        try:
            body_style = styles.add_style('CV Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Arial'
            body_font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(3)
        except:
            body_style = styles['CV Body'] if 'CV Body' in styles else styles['Normal']
        
        try:
            bullet_style = styles.add_style('CV Bullet', WD_STYLE_TYPE.PARAGRAPH)
            bullet_font = bullet_style.font
            bullet_font.name = 'Arial'
            bullet_font.size = Pt(11)
            bullet_style.paragraph_format.left_indent = Inches(0.25)
            bullet_style.paragraph_format.space_after = Pt(3)
        except:
            bullet_style = styles['CV Bullet'] if 'CV Bullet' in styles else styles['Normal']
    
    def _parse_markdown_cv(self, cv_content: str) -> Dict[str, Any]:
        sections = {}
        current_section = None
        current_content = []
        
        lines = cv_content.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            
            if line_stripped.startswith('#'):
                if current_section:
                    sections[current_section] = '\n'.join(current_content)
                
                if line_stripped.startswith('# '):
                    sections['name'] = line_stripped[2:].strip()
                    current_section = None
                    current_content = []
                elif line_stripped.startswith('## '):
                    current_section = line_stripped[3:].strip().lower().replace(' ', '_')
                    current_content = []
            else:
                if current_section and line_stripped:
                    current_content.append(line_stripped)
        
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _add_contact_section(self, contact_text: str, style_profile: StyleProfile) -> None:
        if not contact_text:
            return
            
        contact_para = self.document.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if style_profile.contact_format == "horizontal":
            contact_para.add_run(contact_text).font.size = Pt(11)
        else:
            lines = contact_text.split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    contact_para.add_run('\n')
                
                if line.startswith('**') and line.endswith('**'):
                    clean_line = line.replace('**', '')
                    contact_para.add_run(clean_line).font.size = Pt(11)
                else:
                    contact_para.add_run(line).font.size = Pt(11)
    
    def _add_section_heading(self, heading: str, style_profile: StyleProfile) -> None:
        heading_para = self.document.add_paragraph()
        heading_run = heading_para.add_run(heading.upper())
        heading_run.font.bold = True
        heading_run.font.size = Pt(12)
        heading_para.paragraph_format.space_before = Pt(12)
        heading_para.paragraph_format.space_after = Pt(6)
    
    def _add_career_summary(self, summary_text: str) -> None:
        if not summary_text:
            return
        
        para = self.document.add_paragraph(summary_text)
        para.paragraph_format.space_after = Pt(12)
    
    def _add_skills_section(self, skills_text: str, style_profile: StyleProfile) -> None:
        if not skills_text:
            return
        
        lines = skills_text.split('\n')
        skills_per_row = 3
        current_row_skills = []
        
        for line in lines:
            line = line.strip()
            if line.startswith(('•', '-', '*', '○', '▪')):
                skill = line[1:].strip()
                if skill:
                    current_row_skills.append(skill)
                    
                    if len(current_row_skills) == skills_per_row:
                        para = self.document.add_paragraph()
                        skills_line = f"{style_profile.bullet_style} " + f"  {style_profile.bullet_style} ".join(current_row_skills)
                        para.add_run(skills_line)
                        current_row_skills = []
        
        if current_row_skills:
            para = self.document.add_paragraph()
            skills_line = f"{style_profile.bullet_style} " + f"  {style_profile.bullet_style} ".join(current_row_skills)
            para.add_run(skills_line)
    
    def _add_experience_section(self, experience_text: str, style_profile: StyleProfile) -> None:
        if not experience_text:
            return
        
        lines = experience_text.split('\n')
        current_job = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if '|' in line and not line.startswith(('•', '-', '*', '**')):
                if current_job:
                    self.document.add_paragraph().paragraph_format.space_after = Pt(6)
                
                parts = [p.strip() for p in line.split('|')]
                job_para = self.document.add_paragraph()
                
                if len(parts) >= 2:
                    title_run = job_para.add_run(parts[0])
                    title_run.font.bold = True
                    
                    company_run = job_para.add_run(f" | {parts[1]}")
                    company_run.font.bold = False
                    
                    if len(parts) >= 3:
                        dates_run = job_para.add_run(f" | {parts[2]}")
                        dates_run.font.italic = True
                
                job_para.paragraph_format.space_after = Pt(3)
                current_job = True
            
            elif line.startswith(('•', '-', '*', '○', '▪')) and current_job:
                bullet_text = line[1:].strip()
                if bullet_text:
                    bullet_para = self.document.add_paragraph()
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
                    bullet_para.add_run(f"{style_profile.bullet_style} {bullet_text}")
                    bullet_para.paragraph_format.space_after = Pt(3)
    
    def _add_education_section(self, education_text: str) -> None:
        if not education_text:
            return
        
        lines = education_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and '|' in line:
                edu_para = self.document.add_paragraph()
                parts = [p.strip() for p in line.split('|')]
                
                if len(parts) >= 1:
                    degree_run = edu_para.add_run(parts[0])
                    degree_run.font.bold = True
                
                if len(parts) >= 2:
                    edu_para.add_run(f" | {parts[1]}")
                
                if len(parts) >= 3:
                    date_run = edu_para.add_run(f" | {parts[2]}")
                    date_run.font.italic = True
                
                edu_para.paragraph_format.space_after = Pt(6)
    
    def _add_certifications_section(self, certs_text: str, style_profile: StyleProfile) -> None:
        if not certs_text:
            return
        
        lines = certs_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line.startswith(('•', '-', '*', '○', '▪')):
                cert_text = line[1:].strip()
                if cert_text:
                    cert_para = self.document.add_paragraph()
                    cert_para.add_run(f"{style_profile.bullet_style} {cert_text}")
                    cert_para.paragraph_format.space_after = Pt(3)
            elif line:
                cert_para = self.document.add_paragraph()
                cert_para.add_run(f"{style_profile.bullet_style} {line}")
                cert_para.paragraph_format.space_after = Pt(3)
    
    def export_to_docx(self, cv_content: str, style_profile: StyleProfile, 
                       output_path: str, name: str = None) -> str:
        try:
            self.document = Document()
            
            self._setup_document_styles(style_profile)
            
            sections = self._parse_markdown_cv(cv_content)
            
            if name or sections.get('name'):
                title_para = self.document.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.add_run(name or sections.get('name', 'CV'))
                title_run.font.bold = True
                title_run.font.size = Pt(16)
                title_para.paragraph_format.space_after = Pt(12)
            
            section_order = [
                ('contact_information', 'CONTACT INFORMATION'),
                ('career_summary', 'CAREER SUMMARY'), 
                ('skills', 'SKILLS'),
                ('experience', 'EXPERIENCE'),
                ('education', 'EDUCATION'),
                ('certifications', 'CERTIFICATIONS')
            ]
            
            for section_key, section_title in section_order:
                if section_key in sections:
                    self._add_section_heading(section_title, style_profile)
                    
                    if section_key == 'contact_information':
                        self._add_contact_section(sections[section_key], style_profile)
                    elif section_key == 'career_summary':
                        self._add_career_summary(sections[section_key])
                    elif section_key == 'skills':
                        self._add_skills_section(sections[section_key], style_profile)
                    elif section_key == 'experience':
                        self._add_experience_section(sections[section_key], style_profile)
                    elif section_key == 'education':
                        self._add_education_section(sections[section_key])
                    elif section_key == 'certifications':
                        self._add_certifications_section(sections[section_key], style_profile)
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            self.document.save(output_path)
            
            logger.info(f"CV exported to DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise e
    
    def export_cover_letter_to_docx(self, cover_letter_content: str, 
                                   output_path: str, applicant_name: str = None) -> str:
        try:
            self.document = Document()
            
            if applicant_name:
                header_para = self.document.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                header_run = header_para.add_run(applicant_name)
                header_run.font.bold = True
                header_run.font.size = Pt(12)
                header_para.paragraph_format.space_after = Pt(18)
            
            paragraphs = cover_letter_content.split('\n\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    para = self.document.add_paragraph(para_text)
                    para.paragraph_format.space_after = Pt(12)
                    para.paragraph_format.line_spacing = 1.15
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            self.document.save(output_path)
            
            logger.info(f"Cover letter exported to DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting cover letter to DOCX: {e}")
            raise e

@st.cache_resource 
def get_docx_exporter():
    return DocxExporter()